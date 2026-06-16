import fitz  # PyMuPDF
from groq import Groq
import os
import json
import io
import docx
import base64
import time
from dotenv import load_dotenv

# Load environment variables for independent imports/tests
load_dotenv()

def get_groq_client():
    # Instantiate inside function to prevent crash if key is missing during server boot
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
         raise ValueError("GROQ_API_KEY environment variable is not set")
    return Groq(api_key=api_key)

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_docx(docx_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(docx_bytes))
    paragraphs_text = [p.text for p in doc.paragraphs]
    
    # Also capture table content if present
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs_text.append(cell.text)
                
    return "\n".join(paragraphs_text)

def clean_json_response(result_text: str) -> str:
    # Clean up Markdown code block wrapping if LLM returned it
    result_text = result_text.strip()
    if result_text.startswith("```json"):
        result_text = result_text[7:]
    if result_text.startswith("```"):
        result_text = result_text[3:]
    if result_text.endswith("```"):
        result_text = result_text[:-3]
    return result_text.strip()

def extract_details_from_text(raw_text: str) -> dict:
    prompt = f"""
You are an expert at reading Indian job offer letters.
Extract the following information from this offer letter text.
Return ONLY a valid JSON object with these exact keys:
- role: job title/designation (string)
- company: company name (string)  
- city: work location city (string)
- ctc_lpa: total CTC in LPA as a number (float). Convert if given in rupees per annum.
- experience_years: years of experience required or candidate's experience (int, default 0 for freshers)

If any field cannot be found, use null.
Return only the JSON, no explanation, no markdown.

Offer Letter Text:
{raw_text[:3000]}
"""
    client = get_groq_client()
    last_error = None
    for attempt in range(2):
        try:
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=500,
                temperature=0.1
            )
            result = response.choices[0].message.content.strip()
            result_cleaned = clean_json_response(result)
            return json.loads(result_cleaned)
        except Exception as e:
            last_error = e
            if attempt == 1:
                raise RuntimeError(f"Groq API details extraction failed after 2 attempts: {str(e)}")
            time.sleep(1)

def extract_offer_details_from_image(image_bytes: bytes, filename: str) -> dict:
    base64_image = base64.b64encode(image_bytes).decode('utf-8')
    mime_type = "image/png" if filename.lower().endswith(".png") else "image/jpeg"
    
    prompt = """
You are an expert at reading Indian job offer letters from screenshots and photos.
Analyze this image and extract the following details.
Return ONLY a valid JSON object with these exact keys:
- role: job title/designation (string)
- company: company name (string)  
- city: work location city (string)
- ctc_lpa: total CTC in LPA as a number (float). Convert if given in rupees per annum.
- experience_years: years of experience required or candidate's experience (int, default 0 for freshers)

If any field cannot be found, use null.
Return only the JSON, no explanation, no markdown.
"""
    client = get_groq_client()
    last_error = None
    for attempt in range(2):
        try:
            response = client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=500,
                temperature=0.1
            )
            result = response.choices[0].message.content.strip()
            result_cleaned = clean_json_response(result)
            return json.loads(result_cleaned)
        except Exception as e:
            last_error = e
            if attempt == 1:
                raise RuntimeError(f"Groq API Vision details extraction failed after 2 attempts: {str(e)}")
            time.sleep(1)

def extract_offer_details(file_bytes: bytes, filename: str) -> dict:
    fn_lower = filename.lower()
    if fn_lower.endswith(".pdf"):
        raw_text = extract_text_from_pdf(file_bytes)
        return extract_details_from_text(raw_text)
    elif fn_lower.endswith(".docx"):
        raw_text = extract_text_from_docx(file_bytes)
        return extract_details_from_text(raw_text)
    elif fn_lower.endswith((".jpg", ".jpeg", ".png")):
        return extract_offer_details_from_image(file_bytes, filename)
    else:
        raise ValueError("Unsupported file type. Only PDF, DOCX, and images (JPG/PNG) are accepted.")
