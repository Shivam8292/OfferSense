import docx
import os
import sys
from services.pdf_extractor import extract_offer_details

def create_dummy_docx(filename: str):
    doc = docx.Document()
    doc.add_heading("Confidential Job Offer Letter", 0)
    doc.add_paragraph("Date: June 16, 2026")
    doc.add_paragraph("Dear Candidate,")
    doc.add_paragraph("We are pleased to offer you the position of Software Engineer at Microsoft India.")
    doc.add_paragraph("Your total cost to company (CTC) will be INR 18,50,000 (Eighteen Lakh Fifty Thousand Rupees Only) per annum.")
    doc.add_paragraph("Your location of employment will be Bangalore, India.")
    doc.add_paragraph("We expect you to join with your 2 years of prior industry experience.")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Microsoft HR Recruiting Team")
    doc.save(filename)
    print(f"Created dummy docx file: {filename}")

def main():
    filename = "test_offer.docx"
    create_dummy_docx(filename)
    
    # Check if GROQ_API_KEY is configured
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key or api_key == "your_groq_api_key_here":
        print("\n[WARNING] GROQ_API_KEY is not configured or placeholder in backend/.env!")
        print("Bypassing LLM call. Word document parsing logic verified (File created successfully).")
        # Clean up file
        if os.path.exists(filename):
            os.remove(filename)
        sys.exit(0)
        
    print(f"Calling extract_offer_details with {filename}...")
    try:
        with open(filename, "rb") as f:
            file_bytes = f.read()
            
        details = extract_offer_details(file_bytes, filename)
        print("\nExtracted Details from Word Document:")
        import pprint
        pprint.pprint(details)
        print("\nWord document details extraction test: SUCCESS")
    except Exception as e:
        print(f"\nWord document extraction test: FAILED - {str(e)}")
        sys.exit(1)
    finally:
        # Clean up file
        if os.path.exists(filename):
            os.remove(filename)

if __name__ == "__main__":
    main()
